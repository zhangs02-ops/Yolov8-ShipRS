"""生成 ShipRS 性能对比分析 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cell(cell, text, bold=False, align="center", size=10):
    """设置单元格文本和样式"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }[align]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.bold = bold


def shade_cell(cell, color):
    """给单元格设置背景色"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    shading.append(shd)


def add_table(doc, headers, rows, col_widths=None, highlight_last=False):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=10)
        shade_cell(table.rows[0].cells[i], "D9E2F3")

    # 数据行
    for r_idx, row in enumerate(rows):
        is_last = highlight_last and r_idx == len(rows) - 1
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell(cell, val, bold=is_last, size=10)
            if is_last:
                shade_cell(cell, "FFF2CC")

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading(doc, text, level=2):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_note(doc, text):
    """添加注释段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main():
    doc = Document()

    # ── 标题 ──
    title = doc.add_heading("YOLOv8-ShipRS 船舶遥感小目标检测 — 性能对比分析", level=1)
    for run in title.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # ── 一、实验环境 ──
    add_heading(doc, "一、实验环境")
    env_headers = ["项目", "配置"]
    env_rows = [
        ["GPU", "NVIDIA RTX 3080 20GB"],
        ["操作系统", "Ubuntu 20.04 (WSL2)"],
        ["深度学习框架", "PyTorch 2.0.0 + CUDA 11.8"],
        ["检测框架", "Ultralytics 8.4.21"],
        ["数据集", "Airbus Ship Detection (YOLO格式, 1类)"],
        ["训练参数", "500 epochs, SGD, lr0=0.01, batch=64, imgsz=640, fraction=0.3"],
    ]
    add_table(doc, env_headers, env_rows, col_widths=[4, 12])

    # ── 二、模型说明 ──
    add_heading(doc, "二、模型说明")
    model_headers = ["编号", "模型简称", "配置文件", "改进内容"]
    model_rows = [
        ["1", "YOLOv8（基线）", "yolov8.yaml", "标准 YOLOv8 架构，无改动"],
        ["2", "YOLOv8-AdaptiveDetect", "yolov8-AdaptiveDetect.yaml", "自适应检测头（AdaptiveDetect + FRM）"],
        ["3", "YOLOv8-CA", "yolov8-CA.yaml", "Coordinate Attention 注意力机制（C2f_CA）"],
        ["4", "YOLOv8-P2", "yolov8-P2.yaml", "P2 小目标检测层"],
        ["5", "YOLOv8-ShipRS", "yolov8-ShipRS.yaml", "综合改进（P2 + CA + AdaptiveDetect）"],
    ]
    add_table(doc, model_headers, model_rows, col_widths=[1.5, 4.5, 5, 6], highlight_last=True)

    # ── 三、最终性能对比 ──
    add_heading(doc, "三、最终性能对比（第500轮）")
    perf_headers = ["模型", "Precision", "Recall", "mAP50", "mAP50-95"]
    perf_rows = [
        ["YOLOv8（基线）", "0.8486", "0.7463", "0.8266", "0.5642"],
        ["YOLOv8-AdaptiveDetect", "0.8578", "0.7555", "0.8325", "0.5761"],
        ["YOLOv8-CA", "0.8600", "0.7609", "0.8410", "0.5808"],
        ["YOLOv8-P2", "0.8752", "0.7958", "0.8712", "0.6150"],
        ["YOLOv8-ShipRS", "0.8978", "0.7938", "0.8758", "0.6254"],
    ]
    add_table(doc, perf_headers, perf_rows, col_widths=[4.5, 2.8, 2.8, 2.8, 2.8], highlight_last=True)

    # ── 四、提升幅度 ──
    add_heading(doc, "四、相较于基线模型的提升幅度")
    delta_headers = ["模型", "Precision", "Recall", "mAP50", "mAP50-95"]
    delta_rows = [
        ["AdaptiveDetect", "+0.92pp（+1.1%）", "+0.93pp（+1.2%）", "+0.59pp（+0.7%）", "+1.20pp（+2.1%）"],
        ["CA", "+1.13pp（+1.3%）", "+1.47pp（+2.0%）", "+1.44pp（+1.7%）", "+1.66pp（+2.9%）"],
        ["P2", "+2.65pp（+3.1%）", "+4.96pp（+6.6%）", "+4.46pp（+5.4%）", "+5.09pp（+9.0%）"],
        ["ShipRS", "+4.91pp（+5.8%）", "+4.76pp（+6.4%）", "+4.92pp（+6.0%）", "+6.13pp（+10.9%）"],
    ]
    add_table(doc, delta_headers, delta_rows, col_widths=[4.5, 3, 3, 3, 3], highlight_last=True)
    add_note(doc, "注：pp = percentage points（百分点），括号内为相对提升百分比。")

    # ── 五、模型大小对比 ──
    add_heading(doc, "五、模型大小对比")
    size_headers = ["模型", "权重文件大小", "参数量", "相对基线增长"]
    size_rows = [
        ["YOLOv8（基线）", "6.01 MB", "3,011,043", "—"],
        ["YOLOv8-AdaptiveDetect", "6.21 MB", "~3,143,000", "+3.3%"],
        ["YOLOv8-CA", "6.04 MB", "~3,055,000", "+0.5%"],
        ["YOLOv8-P2", "6.05 MB", "2,926,692", "-2.8%"],
        ["YOLOv8-ShipRS", "6.28 MB", "~3,152,000", "+4.5%"],
    ]
    add_table(doc, size_headers, size_rows, col_widths=[4.5, 3, 3, 3], highlight_last=True)
    add_note(doc, "注：基线和 P2 的参数量为精确值，其余为 FP16 数据量估算值。")

    # ── 六、训练时间对比 ──
    add_heading(doc, "六、训练时间对比")
    time_headers = ["模型", "总训练时间", "平均每轮耗时", "相对基线增长"]
    time_rows = [
        ["YOLOv8（基线）", "24,814s（6.9h）", "~49.6s", "—"],
        ["YOLOv8-AdaptiveDetect", "32,940s（9.1h）", "~65.9s", "+32.7%"],
        ["YOLOv8-CA", "31,163s（8.7h）", "~62.3s", "+25.6%"],
        ["YOLOv8-P2", "37,294s（10.4h）", "~74.6s", "+50.3%"],
        ["YOLOv8-ShipRS", "40,531s（11.3h）", "~81.1s", "+63.4%"],
    ]
    add_table(doc, time_headers, time_rows, col_widths=[4.5, 3.5, 3, 3], highlight_last=True)

    # ── 七、综合性价比分析 ──
    add_heading(doc, "七、综合性价比分析")
    cost_headers = ["模型", "参数增长", "训练时间增长", "mAP50-95提升", "综合评价"]
    cost_rows = [
        ["CA", "+0.5%", "+25.6%", "+1.66pp（+2.9%）", "最轻量改进，几乎零参数开销"],
        ["P2", "-2.8%", "+50.3%", "+5.09pp（+9.0%）", "参数量反而减少，性能大幅提升"],
        ["AdaptiveDetect", "+3.3%", "+32.7%", "+1.20pp（+2.1%）", "提升较温和，性价比一般"],
        ["ShipRS", "+4.5%", "+63.4%", "+6.13pp（+10.9%）", "综合最优，提升最大"],
    ]
    add_table(doc, cost_headers, cost_rows, col_widths=[3.5, 2.5, 3, 3.5, 5], highlight_last=True)

    # ── 保存 ──
    output_path = "/home/zhangs02/yolov8shipRS/ShipRS性能对比分析.docx"
    doc.save(output_path)
    print(f"已保存到: {output_path}")


if __name__ == "__main__":
    main()
